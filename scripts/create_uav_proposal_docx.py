from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


NEW_TITLE = "基于Faster R-ViT与无人机航线规划的玉米雄穗检测系统"

INTRO_APPEND = (
    " 补充无人机规划后，系统进一步形成“田块边界建模—无人机覆盖式航线规划—"
    "低空图像采集—Faster R-ViT雄穗检测—检测结果空间映射—去雄作业路径规划”的闭环。"
    "巡检无人机根据田块边界、飞行高度、相机视场角、航向重叠率和旁向重叠率自动生成往复式航线，"
    "在保证冠层图像完整覆盖的同时记录GPS/北斗定位信息，为后续雄穗位置反算、密度统计和精准去雄提供数据基础。"
)

PURPOSE_APPEND = (
    " 在无人机规划层面，本项目将飞行任务参数化：以田块边界、作物行方向、飞行高度、"
    "相机视场角、拍摄间隔和重叠率作为输入，输出可执行航点序列和复飞策略。"
    "该设计能够减少人工巡田的随机性，使图像采集尺度更稳定、覆盖更完整，并将检测结果由单幅图像扩展到田块空间分布分析。"
)

CONTENT_APPEND = (
    "\n4.无人机巡检航线规划与检测结果空间映射研究。\n"
    "本项目新增无人机规划模块。首先根据玉米制种田边界构建二维作业区域，结合相机视场角、飞行高度和旁向重叠率计算航线间距，"
    "采用往复式覆盖航线完成全田图像采集；其次记录每张图像对应的航点、时间、高度和姿态信息，"
    "将Faster R-ViT输出的雄穗检测框与无人机位置信息关联，形成田块网格级雄穗数量统计和分布热力图；"
    "最后根据未去雄植株的空间分布，为去雄无人机生成优先级作业路径，支持重点区域复查和精准去雄。"
)

ROUTE_APPEND = (
    "\n（7）根据田块边界、飞行高度、相机视场角和重叠率生成巡检无人机覆盖式航线，"
    "推荐采用往复式航线保证图像连续覆盖，并保留起飞点、返航点和异常复飞航点。\n"
    "（8）采集阶段同步记录图像编号、航点坐标、高度、时间和姿态信息，将检测结果映射到田块坐标系，"
    "输出雄穗密度热力图、重点复查区域和去雄作业航点。\n"
    "（9）去雄无人机依据检测热区和单株定位结果进行路径优化，优先处理高密度区域和漏检风险较高区域，"
    "完成后通过复飞图像验证去雄效果。"
)

OUTCOME_APPEND = (
    "\n4.形成无人机巡检航线规划模块，能够根据田块尺寸、飞行高度、相机视场角和重叠率生成覆盖式航点序列，并输出航程、预计飞行时间和航线示意图。\n"
    "5.形成玉米雄穗检测结果空间化表达方法，将单张图像检测框汇总为田块网格级雄穗分布图，为复查和去雄作业提供辅助决策。"
)


def unique_cells(document: Document):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell


def replace_text_everywhere(document: Document, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
    for cell in unique_cells(document):
        for paragraph in cell.paragraphs:
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)


def append_to_cell(document: Document, marker: str, addition: str) -> None:
    for cell in unique_cells(document):
        text = " ".join(cell.text.split())
        if marker in text and addition[:20] not in text:
            cell.add_paragraph(addition)
            return
    raise ValueError(f"marker not found: {marker}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    document = Document(args.input)
    replace_text_everywhere(document, "基于Faster R-ViT的玉米雄穗检测系统", NEW_TITLE)
    append_to_cell(document, "一、项目简介", INTRO_APPEND)
    append_to_cell(document, "二、项目实施的目的、意义", PURPOSE_APPEND)
    append_to_cell(document, "项目研究内容和目标", CONTENT_APPEND)
    append_to_cell(document, "五、项目技术路线", ROUTE_APPEND)
    append_to_cell(document, "六、项目预期成果及说明", OUTCOME_APPEND)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(output)


if __name__ == "__main__":
    main()
